import re
import docx

def simple_text(text):
    text = re.sub(r'(\w+)-\n(\w+)', r'\1\2', text) 
    text = text.replace('\n', ' ')
    text = re.sub(r'\s+', ' ', text)
    return text

def extr_txt(doc):
    return [para.text for para in doc.paragraphs]

def find_opt(text):
    pattern = re.compile(r'\((\b[A-Z]{2,10}\b)\)')
    return [(m.group(1), m.start()) for m in pattern.finditer(text)]

def find_full(acronym, text, idx):
    max_words = min(len(acronym) + 5, 15)
    window = text[max(0, idx - 300):idx].strip()
    words = re.findall(r'\w+', window)[-max_words:]

    acronym_chars = list(acronym.lower())
    for i in range(len(words)):
        seq = words[i:]
        if len(seq) < len(acronym_chars):
            continue
        candidate = seq[:len(acronym_chars)]
        if [w[0].lower() for w in candidate] == acronym_chars:
            return ' '.join(candidate)
    return None

def extract_pairs_from_docx(filename):
    doc = docx.Document(filename)
    raw_blocks = extr_txt(doc)
    full_text = simple_text(' '.join(raw_blocks))
    candidates = find_opt(full_text)

    results = []
    seen = set()

    for acronym, idx in candidates:
        full_form = find_full(acronym, full_text, idx)
        if full_form:
            key = (acronym.lower(), full_form.lower())
            if key not in seen:
                seen.add(key)
                results.append((acronym, full_form))
    return results

def write_to_docx(pairs, output_filename):
    doc = docx.Document()
    doc.add_heading('Acronym List', 1)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Acronym'
    hdr[1].text = 'Full Form'
    for acronym, full_form in pairs:
        row = table.add_row().cells
        row[0].text = acronym
        row[1].text = full_form
    doc.save(output_filename)

pairs = extract_pairs_from_docx('input.docx')
print('b')
if pairs:
    write_to_docx(pairs, 'acronym_output.docx')
else:
    print("No valid acronym-full form pairs found.")
print('a')
